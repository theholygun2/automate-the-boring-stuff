import docx
import readDocx

doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))
#
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].text)
# print(len(doc.paragraphs[1].runs))
# print(doc.paragraphs[1].runs[0].text)
# print(doc.paragraphs[1].runs[1].text)
# print(doc.paragraphs[5].text)
# print(doc.paragraphs[6].text)

#print(readDocx.getText('demo.docx'))

'''Styling Paragraph and Run Objects'''

'''Creating Word Documents with Nondefault Styles'''

'''Run Attributes'''

print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text))
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')
